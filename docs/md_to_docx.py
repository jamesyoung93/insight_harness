"""Build the Insight Harness essay and pilot brief from Markdown.

The builder supports the subset used by the two source documents: headings,
bold/italic spans, external hyperlinks, inline images with accessible alt text,
captions, real Word bullet lists, and Markdown tables.  The essay preserves its
editorial styling; the internal brief uses the decision_memo design preset.

Usage:
    python docs/md_to_docx.py docs/why_we_built_the_insight_harness.md
    python docs/md_to_docx.py docs/trusted_analytics_pilot_brief.md \
        --profile decision_memo
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|\[[^\]]+\]\([^)]+\))")
LINK = re.compile(r"^\[([^\]]+)\]\(([^)]+)\)$")
IMAGE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")
TABLE_RULE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")

TEAL = RGBColor(0x0E, 0x7C, 0x7B)
INK = RGBColor(0x1F, 0x29, 0x37)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x5F, 0x6B, 0x76)
BLACK = RGBColor(0x00, 0x00, 0x00)


def _set_font(run, name: str, size: float | None = None,
              color: RGBColor | None = None, bold: bool | None = None,
              italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _style_font(style, name: str, size: float, color: RGBColor,
                bold: bool | None = None, italic: bool | None = None) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def _add_hyperlink(paragraph, text: str, url: str, profile: str) -> None:
    rid = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial" if profile == "decision_memo" else "Calibri")
    fonts.set(qn("w:hAnsi"), "Arial" if profile == "decision_memo" else "Calibri")
    rpr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5" if profile == "decision_memo" else "0E7C7B")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    node = OxmlElement("w:t")
    node.set(qn("xml:space"), "preserve")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_rich(paragraph, text: str, profile: str) -> None:
    """Add the supported inline Markdown while preserving link relationships."""
    for chunk in INLINE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            paragraph.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            paragraph.add_run(chunk[1:-1]).italic = True
        elif match := LINK.match(chunk):
            _add_hyperlink(paragraph, match.group(1), match.group(2), profile)
        else:
            paragraph.add_run(chunk)


def _set_cell_margins(cell, top: int = 80, start: int = 120,
                      bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start),
                        ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    """Apply fixed 9360-DXA geometry for the decision-memo capability table."""
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        for existing in list(tbl_pr.findall(qn(tag))):
            tbl_pr.remove(existing)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _add_markdown_table(doc: Document, header: list[str], rows: list[list[str]],
                        profile: str) -> None:
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    widths = [1700, 2300, 3200, 2160] if len(header) == 4 else \
        [9360 // len(header)] * len(header)
    widths[-1] += 9360 - sum(widths)

    header_row = table.rows[0]
    tr_pr = header_row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for index, value in enumerate(header):
        cell = header_row.cells[index]
        _shade_cell(cell, "F2F4F7")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_rich(p, value, profile)
        for run in p.runs:
            _set_font(run, "Arial", 9.5, INK, bold=True)

    for row_values in rows:
        row = table.add_row()
        padded = row_values + [""] * (len(header) - len(row_values))
        for index, value in enumerate(padded[:len(header)]):
            p = row.cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.03
            add_rich(p, value, profile)
            for run in p.runs:
                _set_font(run, "Arial", 9.2, INK, bold=(index == 0))

    _set_table_geometry(table, widths)


def _set_alt_text(inline_shape, alt_text: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text[:120])


def _add_image(doc: Document, md_path: Path, alt_text: str, target: str) -> None:
    image_path = (md_path.parent / target).resolve()
    if not image_path.exists():
        raise FileNotFoundError(f"image referenced by Markdown is missing: {image_path}")
    if doc.paragraphs:
        preceding = doc.paragraphs[-1]
        if 0 < len(preceding.text.split()) <= 35:
            preceding.paragraph_format.keep_with_next = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    if image_path.name == "01-trust-gap.png":
        width = 6.25
    elif image_path.name == "06-trust-record.png":
        width = 4.6
    else:
        width = 5.75
    shape = run.add_picture(str(image_path), width=Inches(width))
    _set_alt_text(shape, alt_text)


def _add_page_number(paragraph) -> None:
    paragraph.add_run("Page ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    fallback = OxmlElement("w:t")
    fallback.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = OxmlElement("w:r")
    run.extend([begin, instr, separate, fallback, end])
    paragraph._p.append(run)


def _configure_page(doc: Document) -> None:
    doc.settings.odd_and_even_pages_header_footer = False
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = False


def _remove_style_borders(style) -> None:
    p_pr = style._element.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is not None:
        p_pr.remove(borders)


def _configure_essay(doc: Document) -> None:
    _configure_page(doc)
    normal = doc.styles["Normal"]
    _style_font(normal, "Calibri", 11, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12

    title = doc.styles["Title"]
    _style_font(title, "Calibri", 26, INK, bold=True)
    _remove_style_borders(title)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.styles["Subtitle"]
    _style_font(subtitle, "Calibri", 14, TEAL, bold=False)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(10)

    h2 = doc.styles["Heading 2"]
    _style_font(h2, "Calibri", 15, TEAL, bold=True)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    _style_font(h3, "Calibri", 12, TEAL, bold=True)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    _style_font(caption, "Calibri", 9.5, MUTED, italic=True)
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(9)
    caption.paragraph_format.keep_together = True

    bullet = doc.styles["List Bullet"]
    _style_font(bullet, "Calibri", 11, INK)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(6)
    bullet.paragraph_format.line_spacing = 1.15


def _configure_decision_memo(doc: Document) -> None:
    _configure_page(doc)
    normal = doc.styles["Normal"]
    _style_font(normal, "Arial", 11, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    title = doc.styles["Title"]
    _style_font(title, "Arial", 23, BLACK, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)

    h2 = doc.styles["Heading 2"]
    _style_font(h2, "Arial", 13, BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(7)
    h2.paragraph_format.space_after = Pt(3)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    _style_font(h3, "Arial", 12, DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(7)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.keep_with_next = True

    bullet = doc.styles["List Bullet"]
    _style_font(bullet, "Arial", 11, INK)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(8)
    bullet.paragraph_format.line_spacing = 1.167

    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = header.add_run("TRUSTED ANALYTICS PILOT")
    _set_font(left, "Arial", 8.5, MUTED, bold=True)
    header.add_run("\t")
    right = header.add_run("INTERNAL WORKING BRIEF")
    _set_font(right, "Arial", 8.5, MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    _add_page_number(footer)
    for run in footer.runs:
        _set_font(run, "Arial", 8.5, MUTED)


def _add_memo_title(doc: Document, title_text: str) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("PILOT ALIGNMENT BRIEF")
    _set_font(run, "Arial", 9, TEAL, bold=True)

    title = doc.add_paragraph(style="Title")
    title.add_run(title_text)
    title.paragraph_format.keep_with_next = True

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(8)
    meta.paragraph_format.keep_with_next = True
    run = meta.add_run("Purpose: pilot design alignment  |  Status: decision brief")
    _set_font(run, "Arial", 9.5, MUTED)


def convert(md_path: Path, profile: str | None = None,
            out_path: Path | None = None) -> Path:
    md_path = md_path.resolve()
    lines = md_path.read_text(encoding="utf-8").splitlines()
    profile = profile or ("decision_memo" if "pilot_brief" in md_path.stem else "essay")
    if profile not in {"essay", "decision_memo"}:
        raise ValueError(f"unknown profile: {profile}")

    doc = Document()
    if profile == "decision_memo":
        _configure_decision_memo(doc)
    else:
        _configure_essay(doc)

    index = 0
    standfirst_seen = False
    expect_caption = False
    title_text = ""
    while index < len(lines):
        line = lines[index].rstrip()
        if not line.strip():
            index += 1
            continue

        if line.startswith("|") and index + 1 < len(lines) and TABLE_RULE.match(lines[index + 1]):
            header = _split_table_row(line)
            rows: list[list[str]] = []
            index += 2
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                rows.append(_split_table_row(lines[index]))
                index += 1
            _add_markdown_table(doc, header, rows, profile)
            continue

        if match := IMAGE.match(line):
            _add_image(doc, md_path, match.group(1), match.group(2))
            expect_caption = True
            index += 1
            continue

        if line.startswith("# "):
            title_text = line[2:].strip()
            if profile == "decision_memo":
                _add_memo_title(doc, title_text)
            else:
                title = doc.add_paragraph(style="Title")
                title.add_run(title_text)
                title.paragraph_format.keep_with_next = True
            index += 1
            continue

        if line.startswith("### "):
            if profile == "essay" and not standfirst_seen:
                p = doc.add_paragraph(line[4:], style="Subtitle")
                p.paragraph_format.keep_with_next = True
            else:
                doc.add_paragraph(line[4:], style="Heading 3")
            index += 1
            continue

        if line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 2")
            index += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich(p, line[2:], profile)
            if index + 1 < len(lines) and lines[index + 1].lstrip().startswith("- "):
                p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True
            index += 1
            continue

        if line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            text = line[1:-1]
            if expect_caption:
                p = doc.add_paragraph(style="Caption")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_rich(p, text, profile)
                expect_caption = False
            elif profile == "essay" and not standfirst_seen:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(18)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(text)
                _set_font(run, "Calibri", 12.5, RGBColor(0x33, 0x33, 0x33), italic=True)
                standfirst_seen = True
            else:
                p = doc.add_paragraph()
                add_rich(p, line, profile)
            index += 1
            continue

        p = doc.add_paragraph()
        add_rich(p, line, profile)
        index += 1

    doc.core_properties.title = title_text
    doc.core_properties.subject = ("Public practitioner essay" if profile == "essay"
                                   else "Internal pilot alignment brief")
    out = (out_path or md_path.with_suffix(".docx")).resolve()
    doc.save(out)
    return out


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown", nargs="?", default="docs/why_we_built_the_insight_harness.md")
    parser.add_argument("--profile", choices=("essay", "decision_memo"))
    parser.add_argument("--output")
    args = parser.parse_args()
    result = convert(Path(args.markdown), args.profile,
                     Path(args.output) if args.output else None)
    print(f"wrote {result}")


if __name__ == "__main__":
    main()
